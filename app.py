import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
from collections import Counter
from io import BytesIO
from fpdf import FPDF
from docx import Document
import speech_recognition as sr
import tempfile
import os

def limpiar_y_tokenizar(texto):
    texto = texto.lower()
    caracteres_invalidos = ".,;:\'\"!?¿¡\n"
    for c in caracteres_invalidos:
        texto = texto.replace(c, ' ')
    palabras = texto.split()
    return palabras

def obtener_frecuencias(palabras, top_n=10):
    contador = Counter(palabras)
    mas_comunes = contador.most_common(top_n)
    menos_comunes = sorted([item for item in contador.items() if item[1] <= 2], key=lambda x: x[1])[:top_n]
    return mas_comunes, menos_comunes

def generar_senal_con_palabras(mas, menos):
    total = len(mas) + len(menos)
    puntos_por_etiqueta = 20
    x = np.linspace(0, 2 * np.pi, total * puntos_por_etiqueta)
    senal = np.sin(x)
    etiquetas = [''] * len(senal)
    for i in range(total):
        idx = i * puntos_por_etiqueta
        if i < len(mas):
            etiquetas[idx] = mas[i][0]
        else:
            etiquetas[idx] = menos[i - len(mas)][0]
    return x, senal, etiquetas, puntos_por_etiqueta

def exportar_pdf(fig):
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpfile:
        fig.savefig(tmpfile.name)
        tmpfile.flush()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Análisis de Espectro de Texto", ln=True, align='C')
        pdf.image(tmpfile.name, x=10, y=30, w=180)
        output = BytesIO()
        pdf.output(output)
        output.seek(0)
    return output

def exportar_word(mas, menos):
    doc = Document()
    doc.add_heading("Análisis de Espectro de Texto", 0)
    doc.add_heading("Palabras más frecuentes", level=1)
    for palabra, freq in mas:
        doc.add_paragraph(f"{palabra}: {freq}")
    doc.add_heading("Palabras menos frecuentes", level=1)
    for palabra, freq in menos:
        doc.add_paragraph(f"{palabra}: {freq}")
    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output

def transcribir_audio(audio_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
        tmp.write(audio_bytes)
        tmp_path = tmp.name
    recognizer = sr.Recognizer()
    with sr.AudioFile(tmp_path) as source:
        audio_data = recognizer.record(source)
        try:
            texto = recognizer.recognize_google(audio_data, language="es-ES")
            return texto
        except sr.UnknownValueError:
            return ""
        except sr.RequestError:
            return "Error de conexión con el servicio de reconocimiento de voz."

# Interfaz Streamlit
st.set_page_config(layout="wide")
st.title("📊 Analizador de Espectro de Texto con Senoides")

modo = st.radio("Selecciona el modo de entrada:", ["Texto manual", "Subir audio (voz)"])

texto = ""

if modo == "Texto manual":
    texto = st.text_area("Ingresa o pega tu texto aquí:", height=300)
else:
    audio_file = st.file_uploader("Sube un archivo de audio (.wav):", type=["wav"])
    if audio_file is not None:
        with st.spinner("Transcribiendo audio..."):
            texto = transcribir_audio(audio_file.read())
        st.success("Transcripción completada:")
        st.write(texto)

if st.button("🔍 Analizar"):
    if texto.strip() == "":
        st.warning("Por favor, ingresa texto o sube un audio válido.")
    else:
        palabras = limpiar_y_tokenizar(texto)
        mas, menos = obtener_frecuencias(palabras)
        x, senal, etiquetas, step = generar_senal_con_palabras(mas, menos)

        # Mostrar listas visuales
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("🔝 Palabras más frecuentes")
            for palabra, freq in mas:
                st.write(f"{palabra}: {freq}")
        with col2:
            st.subheader("🔻 Palabras menos frecuentes")
            for palabra, freq in menos:
                st.write(f"{palabra}: {freq}")

        # Gráfica senoidal
        st.subheader("📈 Gráfica Senoidal de Frecuencia de Palabras")
        fig, ax = plt.subplots(figsize=(14, 6))
        ax.plot(x, senal)
        for i in range(0, len(etiquetas), step):
            if etiquetas[i]:
                ax.text(x[i], senal[i], etiquetas[i], fontsize=9, ha='center',
                        va='bottom' if senal[i] > 0 else 'top')
        ax.set_title("Palabras más (crestas) y menos (valles) frecuentes sobre señal senoidal")
        ax.set_xlabel("Tiempo (arbitrario)")
        ax.set_ylabel("Amplitud")
        ax.grid(True)
        st.pyplot(fig)

        # Análisis FFT del espectro
        st.subheader("🎼 Espectro de Frecuencia (FFT)")
        senal_fft = np.abs(np.fft.fft(senal))
        fig_fft, ax_fft = plt.subplots(figsize=(14, 4))
        ax_fft.plot(senal_fft)
        ax_fft.set_title("Transformada de Fourier del Espectro de Palabras")
        ax_fft.set_xlabel("Frecuencia")
        ax_fft.set_ylabel("Magnitud")
        ax_fft.grid(True)
        st.pyplot(fig_fft)

        # Descargas
        col1, col2 = st.columns(2)
        with col1:
            pdf_file = exportar_pdf(fig)
            st.download_button("📄 Descargar PDF", data=pdf_file, file_name="espectro_texto.pdf")
        with col2:
            word_file = exportar_word(mas, menos)
            st.download_button("📝 Descargar Word", data=word_file, file_name="frecuencias_texto.docx")
